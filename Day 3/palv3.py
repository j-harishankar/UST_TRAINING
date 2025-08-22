import tkinter as tk
from tkinter import ttk, messagebox
import os, csv, shutil
from docx import Document
from tkinter import filedialog

# File names
txt_file = "Paliendrome.txt"
csv_file = "ReversedTable.csv"
doc_file = "PaliendromeDoc.docx"

# Root window
root = tk.Tk()
root.geometry("440x400")
root.title('Palindrome Check')
root.resizable(True, True)

# Ensure files exist
if not os.path.exists(txt_file):
    with open(txt_file, 'w') as fp:
        fp.write("PALIENDROME LIST\n===================\n")

if not os.path.exists(doc_file):
    document = Document()
    document.add_heading('PALIENDROME CHECK', 0)
    document.save(doc_file)

if not os.path.exists(csv_file):
    with open(csv_file, 'w', newline='') as fp:
        writer = csv.writer(fp)
        writer.writerow(['Given', 'Reversed'])

# Clear input and output
def clear_all():
    given_string_Entry.delete(0, tk.END)
    reverse_value_label.config(text="")

# Verify input is alphanumeric
def verify_user_entry(event):
    reverse_value_label.config(text="")
    if not event.char.isalnum() and event.keysym not in ("BackSpace", "Delete"):
        given_string_Entry.delete(len(given_string_Entry.get()) - 1, tk.END)

# Privilege mode
privilege = tk.IntVar()

# Reverse string and check palindrome
def show_reverse():
    given_str = given_string_Entry.get()
    if not given_str:
        messagebox.showwarning("Empty", "Input is mandatory")
        return

    result_str = given_str[::-1]
    reverse_value_label.config(text=result_str)

    record_data(privilege.get(), given_str, result_str)

    if given_str == result_str:
        messagebox.showinfo(result_str, "Palindrome")
    else:
        messagebox.showinfo(result_str, "Not a Palindrome")

# Record data based on admin/user mode
def record_data(user_type, given_str, reversed_str):
    if user_type == 1:  # Admin only
        if given_str == reversed_str:
            # Save to TXT
            with open(txt_file, 'a') as fp:
                fp.write(given_str + '\n')

            # Save to DOC (avoid duplicates)
            document = Document(doc_file)
            existing_texts = [p.text for p in document.paragraphs]
            if given_str not in existing_texts:
                document.add_paragraph(given_str)
                document.save(doc_file)
        else:
            # Save to CSV
            with open(csv_file, 'a', newline='') as fp:
                writer = csv.writer(fp)
                writer.writerow([given_str, reversed_str])

# Download Doc
def download_doc():
    save_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")]
    )
    if save_path:
        shutil.copyfile(doc_file, save_path)
        messagebox.showinfo("Success", f"Document saved to:\n{save_path}")

# Download CSV
def download_csv():
    save_path = filedialog.asksaveasfilename(
        defaultextension=".csv",
        filetypes=[("CSV File", "*.csv"), ("All Files", "*.*")]
    )
    if save_path:
        shutil.copyfile(csv_file, save_path)
        messagebox.showinfo("Success", f"CSV saved to:\n{save_path}")

# Widgets
pal_label = ttk.Label(root, text="PALIENDROME CHECK")
pal_label.grid(column=0, row=0, columnspan=3, pady=10)

admin_rdo_btn = ttk.Radiobutton(root, text="ADMINISTRATOR", variable=privilege, value=1, command=clear_all)
admin_rdo_btn.grid(column=0, row=1, sticky=tk.W, padx=20, pady=5)

user_rdo_btn = ttk.Radiobutton(root, text="USER", variable=privilege, value=0)
user_rdo_btn.grid(column=1, row=1, sticky=tk.W, padx=20, pady=5)

given_string_label = ttk.Label(root, text="Given String")
given_string_label.grid(column=0, row=2, sticky=tk.W, padx=5, pady=5)

given_string_Entry = ttk.Entry(root)
given_string_Entry.grid(column=1, row=2, sticky=tk.W, padx=5, pady=5)
given_string_Entry.bind('<KeyRelease>', verify_user_entry)

reverse_btn = ttk.Button(root, text="Reverse", command=show_reverse)
reverse_btn.grid(column=2, row=2, sticky=tk.W, padx=5, pady=5)

clear_btn = ttk.Button(root, text="Clear", command=clear_all)
clear_btn.grid(column=2, row=3, sticky=tk.W, padx=5, pady=5)

reverse_label = ttk.Label(root, text="Reversed String:")
reverse_label.grid(column=0, row=3, sticky=tk.W, padx=10, pady=10)

reverse_value_label = ttk.Label(root, text="", font='Helvetica 12 bold')
reverse_value_label.grid(column=1, row=3, sticky=tk.W, padx=5, pady=5)

download_label = ttk.Label(root, text="DOWNLOAD")
download_label.grid(column=0, row=4, sticky=tk.W, padx=5, pady=5)

download_doc_btn = ttk.Button(root, text="Download Doc", command=download_doc)
download_doc_btn.grid(column=1, row=4, sticky=tk.W, padx=5, pady=5)

download_csv_btn = ttk.Button(root, text="Download CSV", command=download_csv)
download_csv_btn.grid(column=2, row=4, sticky=tk.W, padx=5, pady=5)

root.mainloop()
